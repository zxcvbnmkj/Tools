import argparse
import pypandoc
import os
from docx.shared import RGBColor
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def setup_template_styles(temp_name):
    doc = Document()
    # 修改正文样式
    style_normal = doc.styles['Normal']
    # 设置英文字体
    style_normal.font.name = 'Times New Roman'
    # 设置中文字体
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style_normal.font.size = Pt(12)

    # 设置标题的样式
    heading_configs = {
        'Heading 1': Pt(18),
        'Heading 2': Pt(16),
        'Heading 3': Pt(14),
        'Heading 4': Pt(12),
        'Heading 5': Pt(11),
        'Heading 6': Pt(10.5),
    }
    for style_name, font_size in heading_configs.items():
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.name = '微软雅黑'
            h._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            h.font.size = font_size
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)  # 统一设为黑色
    # 添加一个空段落确保文件结构完整，避免 Word 报错
    doc.add_paragraph("")
    # 保存前清理所有物理段落（只留样式）
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    print(f"创建了新的 word 模板{temp_name}")
    doc.save(temp_name)


def convert_md_to_docx(input_file, use_temp):
    # 输出文件名使用原文件名更换后缀
    output_file = os.path.splitext(input_file)[0] + '.docx'
    # 设置 pandox 的参数
    extra_args = []
    if use_temp and os.path.exists(use_temp):
        # 使用 --reference-doc 指定模板
        extra_args.append(f'--reference-doc={use_temp}')
        print(f"[*] 正在使用模板: {use_temp}")
    else:
        print("[!] 未指定模板或模板不存在，将使用默认样式转换")
    pypandoc.convert_file(input_file, 'docx', outputfile=output_file, extra_args=extra_args)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Markdown 转 Word 工具")
    parser.add_argument("--input_file", default='README.md', help="输入的 Markdown 文件路径")
    parser.add_argument("--not_build_temp", action="store_true", help="是否根据当前样式重置/创建模板")
    parser.add_argument("--temp_name", default="template.docx" , help="新 word 模板的文件名，如果不为空则创建它。无需创建则置为 None")
    parser.add_argument("--use_temp", default="template.docx" , help="使用哪一个 word 模板")
    args = parser.parse_args()
    if not args.not_build_temp:
        setup_template_styles(args.temp_name)
    convert_md_to_docx(args.input_file, args.use_temp)